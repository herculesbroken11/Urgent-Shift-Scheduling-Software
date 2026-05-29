#!/usr/bin/env python3
"""Generate BlueThread Solutions Technical Audit Report as Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "BlueThread_Solutions_Technical_Audit_Report.docx"


def set_cell_shading(cell, fill_hex="E8EEF7"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "2E5090")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_mono_block(doc, text):
    """ASCII diagram / code block in monospace."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()


def add_detail_table(doc, rows):
    """Two-column attribute/detail table for edge function audits."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (attr, detail) in enumerate(rows):
        table.rows[i].cells[0].text = attr
        table.rows[i].cells[1].text = detail
        for ci in range(2):
            for p in table.rows[i].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        set_cell_shading(table.rows[i].cells[0], "E8EEF7")
    doc.add_paragraph()


ARCHITECTURE_DIAGRAM = r"""
+-----------------------------------------------------------------------------+
|                         USER BROWSER (SPA)                                   |
|  React 18 + Vite 5 + TypeScript | React Router | TanStack Query              |
|  src/main.tsx -> App.tsx | DemoContext (optional) | data-adapter hooks       |
+-----------------------------------+-----------------------------------------+
                                    | HTTPS
                    VITE_SUPABASE_URL + VITE_SUPABASE_PUBLISHABLE_KEY (anon)
                    Authorization: Bearer <user JWT>  |  functions.invoke()
                                    v
+-----------------------------------------------------------------------------+
|                         SUPABASE PLATFORM                                    |
|  +----------------+  +----------------------+  +------------------------+ |
|  | Supabase Auth  |  | PostgreSQL + RLS     |  | Edge Functions (Deno)  | |
|  | email/OAuth    |  | PostgREST + 29 RPCs  |  | 13 deployed functions  | |
|  | sessions/JWT   |  | Realtime (messages)  |  | service role after auth| |
|  +----------------+  +----------------------+  +------------------------+ |
|  | pg_cron + pg_net -> process-reminders (scheduled)                         |
+-------+--------------+--------------+--------------+--------------+-----------+
        |              |              |              |              |
        v              v              v              v              v
   +---------+   +----------+   +-----------+   +--------+   +-------------+
   | Resend  |   | Twilio   |   | Google    |   | Intuit |   | Lovable     |
   | (email) |   | (SMS)    |   | Calendar  |   | QBO    |   | Cloud Auth  |
   +---------+   +----------+   +-----------+   +--------+   +-------------+
"""

DATA_FLOW_DIAGRAM = r"""
  [Login/Signup] --> Supabase Auth --> JWT in localStorage
        |
        v
  [AuthContext] loads profiles + user_roles (PostgREST, RLS-scoped)
        |
        +--> [Pages/Hooks] --useAdaptedQuery--> PostgREST / RPC (production)
        |         |
        |         +--> isDemoMode? --> DemoDataContext (in-memory, no DB)
        |
        +--> [Mutations] --> supabase.from().insert/update (RLS enforced)
        |
        +--> [Privileged ops] --> functions.invoke() --> Edge Function
                                      |
                                      +--> authenticateCaller (JWT)
                                      +--> service-role client
                                      +--> Resend / Twilio / Google / QBO APIs
"""

ER_RELATIONSHIPS = [
    ["agencies", "1:N", "profiles", "Tenant root; profiles.agency_id"],
    ["agencies", "1:N", "user_roles", "Role per user per agency"],
    ["agencies", "1:N", "customers", "Requesting organizations"],
    ["customers", "1:N", "locations", "Sites under customer"],
    ["agencies", "1:N", "appointments", "Core scheduling entity"],
    ["customers", "N:1", "appointments", "appointments.customer_id"],
    ["locations", "N:1", "appointments", "appointments.location_id"],
    ["profiles", "N:1", "appointments (interpreter)", "appointments.interpreter_id"],
    ["profiles", "N:1", "appointments (requester)", "appointments.requester_id"],
    ["languages", "N:1", "appointments", "appointments.language_id"],
    ["profiles", "1:N", "interpreter_availability", "Availability windows"],
    ["profiles", "N:M", "languages", "via interpreter_languages"],
    ["agencies", "1:N", "billing_rates", "Customer or default rates"],
    ["customers", "1:N", "invoices", "Agency billing"],
    ["invoices", "1:N", "invoice_line_items", "Line items per invoice"],
    ["agencies", "1:N", "conversations", "Messaging threads"],
    ["conversations", "1:N", "messages", "Message bodies"],
    ["conversations", "N:M", "profiles", "via conversation_participants"],
    ["agencies", "1:N", "google_calendar_connections", "Per-user GCal OAuth"],
    ["agencies", "1:N", "qbo_connections", "QuickBooks OAuth per agency"],
    ["agencies", "1:N", "import_batches", "CSV import runs"],
    ["import_batches", "1:N", "import_batch_rows", "Row-level import results"],
    ["appointments", "1:N", "appointment_history", "Audit trail"],
    ["appointments", "N:1", "appointments", "parent_recurring_id (series)"],
    ["auth.users", "1:1", "profiles", "profiles.id references auth.users"],
    ["platform_roles", "N:1", "auth.users", "Platform owner (global)"],
]

EDGE_FUNCTION_AUDITS = [
    {
        "name": "accept-invitation",
        "path": "supabase/functions/accept-invitation/index.ts",
        "endpoint": "POST /functions/v1/accept-invitation",
        "purpose": "After signup/login, links authenticated user to pending invitation by email match; assigns agency, role, profile updates.",
        "auth": "Yes — Bearer JWT via getClaims(); email from token (not client body).",
        "roles": "Any authenticated user (self-service onboarding).",
        "validation": "Matches invitations by ilike email, status=pending, expires_at > now; returns reason if revoked/expired.",
        "tables": "invitations, profiles, user_roles, customer_requestors (requester), agencies; auth.admin as needed.",
        "errors": "401 missing/invalid token; structured JSON {found, accepted, reason}.",
        "security": "Does not trust client agency_id; server derives from invitation. Service role for writes.",
        "production": "Partial — production-ready pattern; verify invitation email uniqueness and race on concurrent accept.",
    },
    {
        "name": "auth-email-hook",
        "path": "supabase/functions/auth-email-hook/index.ts",
        "endpoint": "POST /functions/v1/auth-email-hook (webhook + preview)",
        "purpose": "Supabase/Lovable auth email hook: renders React Email templates (signup, invite, recovery, magic link, etc.) and sends via Lovable email service.",
        "auth": "Webhook: Lovable signature verification (LOVABLE_API_KEY + verifyWebhookRequest). Preview: Bearer LOVABLE_API_KEY only.",
        "roles": "System webhook (not end-user); no agency role check.",
        "validation": "Webhook payload parser; template type whitelist; HTML escape in branded templates.",
        "tables": "None directly (email delivery external).",
        "errors": "401 invalid signature; 400 invalid payload; 500 missing LOVABLE_API_KEY.",
        "security": "verify_jwt=false in config.toml; relies on webhook HMAC. Do not expose preview endpoint publicly.",
        "production": "Partial — depends on Lovable email infrastructure and secrets in Supabase.",
    },
    {
        "name": "google-calendar-sync",
        "path": "supabase/functions/google-calendar-sync/index.ts",
        "endpoint": "POST /functions/v1/google-calendar-sync",
        "purpose": "Google Calendar OAuth and sync: oauth-callback, status, disconnect, update-settings, sync, delete, bulk-sync appointment events.",
        "auth": "Yes — authenticateCaller (JWT + agency from profile).",
        "roles": "Any authenticated agency member with valid JWT (actions typically admin-driven in UI).",
        "validation": "Per-action body checks (code, redirect_uri, appointment_id, etc.); server-side OAuth secret.",
        "tables": "google_calendar_connections, appointments (gcal_event_id, sync fields), profiles, agencies.",
        "errors": "errorResponse / AuthError pattern; token refresh failures stored on connection.",
        "security": "OAuth tokens stored in DB (service role access). Appointment data sent to Google — third-party disclosure.",
        "production": "Partial — requires GOOGLE_CALENDAR_CLIENT_ID/SECRET; needs production OAuth consent screen.",
    },
    {
        "name": "invite-user",
        "path": "supabase/functions/invite-user/index.ts",
        "endpoint": "POST /functions/v1/invite-user",
        "purpose": "Invite/create/resend/revoke users for roles interpreter, requester, scheduler; uses auth.admin for email lookup.",
        "auth": "Yes — authenticateCaller.",
        "roles": "invite/create: agency_admin only. resend/revoke: agency_admin or scheduler.",
        "validation": "Email required; role whitelist; customer_id validated for requesters; paginated auth.admin.listUsers.",
        "tables": "invitations, profiles, user_roles, customers, customer_requestors; auth.users via admin API.",
        "errors": "403 forbidden role; 400 validation; idempotent upserts documented in code.",
        "security": "Strong server-side role checks; does not use profiles.email alone for auth user lookup.",
        "production": "Partial — production-capable; test invite expiry and email deliverability.",
    },
    {
        "name": "manage-join-request",
        "path": "supabase/functions/manage-join-request/index.ts",
        "endpoint": "POST /functions/v1/manage-join-request",
        "purpose": "Self-join workflow: submit join request; approve/reject by agency_admin.",
        "auth": "Yes — JWT getClaims for all modes.",
        "roles": "submit: authenticated user without agency. approve/reject: agency_admin of same agency.",
        "validation": "Agency active; allow_self_join setting; role in self_join_roles; duplicate request prevention.",
        "tables": "join_requests, agencies, profiles, user_roles, customers (requester validation).",
        "errors": "400/403/404 with jsonError messages.",
        "security": "Agency settings gate self-join; revalidation on approve.",
        "production": "Partial — verify abuse (spam join requests) and rate limiting not present.",
    },
    {
        "name": "platform-admin",
        "path": "supabase/functions/platform-admin/index.ts",
        "endpoint": "POST /functions/v1/platform-admin (JSON body: action)",
        "purpose": "Platform owner console API: bootstrap, seed_by_email, promote/demote owners, agency.update, billing_config.*, usage, user admin, support sessions, etc.",
        "auth": "bootstrap: any valid JWT if no owners exist. seed_by_email: service role key only. Else: requirePlatformOwner.",
        "roles": "platform_owner for most actions; bootstrap is special case.",
        "validation": "Allowed field whitelist on agency.update; cannot demote last owner.",
        "tables": "platform_roles, agencies, platform_audit_log, platform_billing_config, profiles, user_roles, notifications (approval email).",
        "errors": "withCors + jsonError; audit log on mutations.",
        "security": "CRITICAL: bootstrap allows first user to become platform owner. seed_by_email requires service key.",
        "production": "Not production-safe until bootstrap disabled post-seed.",
    },
    {
        "name": "platform-qbo",
        "path": "supabase/functions/platform-qbo/index.ts",
        "endpoint": "POST/GET /functions/v1/platform-qbo",
        "purpose": "Platform-level QuickBooks: OAuth initiate/callback, generate/list/sync platform invoices, bulk issue/sync, orphans, delete.",
        "auth": "callback: OAuth redirect. Other actions: authenticatePlatformOwner.",
        "roles": "platform_owner only.",
        "validation": "Action switch; bulk operation locks; QBO token refresh.",
        "tables": "platform_qbo_connection, platform_invoices, platform_invoice_line_items, platform_qbo_sync_log, agencies.",
        "errors": "AuthError, QBO API errors logged to sync log.",
        "security": "Stores QBO tokens at platform level; high privilege.",
        "production": "Partial — sandbox vs production QBO_ENVIRONMENT must be configured.",
    },
    {
        "name": "process-import",
        "path": "supabase/functions/process-import/index.ts",
        "endpoint": "POST /functions/v1/process-import",
        "purpose": "CSV import pipeline: dry_run, execute (chunked), resume, rollback for customers/locations/appointments/profiles.",
        "auth": "Yes — authenticateCaller.",
        "roles": "agency_admin only.",
        "validation": "VALID_STATUSES, VALID_MODALITIES, header signatures, row-level validation messages, concurrency checks via RPC.",
        "tables": "import_batches, import_batch_rows, customers, locations, appointments, profiles (+ staging flags).",
        "errors": "Structured row results; rollback with dependency checks.",
        "security": "Large CSV attack surface; admin-only reduces risk. Service role for bulk writes.",
        "production": "Partial — requires load testing and staged import review on production data.",
    },
    {
        "name": "process-reminders",
        "path": "supabase/functions/process-reminders/index.ts",
        "endpoint": "POST /functions/v1/process-reminders (cron + manual)",
        "purpose": "Scheduled appointment reminders (24h, 2h, 15m) via email/SMS; cancel_stale_reminders RPC first.",
        "auth": "Service role token OR agency_admin JWT. NO header: allowed (cron backward compat).",
        "roles": "Cron: unauthenticated if no header. Manual: agency_admin.",
        "validation": "Time windows per reminder type; dedup keys; agency/interpreter preference resolution.",
        "tables": "appointments, profiles, agencies, interpreter_notification_prefs, notification_log.",
        "errors": "Returns sent/skipped/errors counts; logs to console.",
        "security": "CRITICAL: unauthenticated invocation path. Must fix before production.",
        "production": "Not production-ready until auth required on all invocations.",
    },
    {
        "name": "qbo-auth",
        "path": "supabase/functions/qbo-auth/index.ts",
        "endpoint": "POST/GET /functions/v1/qbo-auth",
        "purpose": "Agency QuickBooks OAuth and sync: initiate, callback, status, disconnect, sync-appointment, bulk-sync, reconcile, validate-mappings, fetch-qbo-items.",
        "auth": "callback: OAuth (GET). Other: authenticateCaller + agency_admin.",
        "roles": "agency_admin for management actions.",
        "validation": "OAuth state parameter; token refresh; QBO API error handling.",
        "tables": "qbo_connections, qbo_sync_jobs, qbo_sync_log, qbo_item_mappings, customers, appointments, invoices.",
        "errors": "Redirects with qbo_error query params on callback failure.",
        "security": "Stores QBO access/refresh tokens in qbo_connections. Admin-only mutations.",
        "production": "Partial — requires QBO app review for production; webhook separate function.",
    },
    {
        "name": "qbo-webhook",
        "path": "supabase/functions/qbo-webhook/index.ts",
        "endpoint": "POST/GET /functions/v1/qbo-webhook",
        "purpose": "Intuit webhook receiver: validates HMAC signature, maps realmId to agency, processes entity change events idempotently.",
        "auth": "HMAC intuit-signature vs QBO_WEBHOOK_VERIFIER_TOKEN. GET returns OK (validation ping).",
        "roles": "N/A (Intuit system caller).",
        "validation": "Idempotency key per entity operation; skips unknown realmId.",
        "tables": "qbo_webhook_events, qbo_connections.",
        "errors": "401 invalid signature; continues on unknown realm with warning.",
        "security": "Good webhook pattern. Ensure verifier token is strong and rotated.",
        "production": "Partial — needs Intuit webhook subscription configured to production URL.",
    },
    {
        "name": "send-notification",
        "path": "supabase/functions/send-notification/index.ts",
        "endpoint": "POST /functions/v1/send-notification",
        "purpose": "Send in_app, email, or SMS notifications; optional template_id with variable substitution.",
        "auth": "Yes — authenticateCaller.",
        "roles": "agency_admin/scheduler: all channels. interpreter/requester: in_app only.",
        "validation": "Template must belong to agency; target user same agency for in_app; recipient required for email/SMS.",
        "tables": "notification_templates, notifications, notification_log (via deliverAndLog).",
        "errors": "403 cross-agency target; 404 template; delivery failures returned in JSON.",
        "security": "Prevents cross-agency in_app targeting. Email/SMS content may include PHI — minimize.",
        "production": "Partial — depends on Resend/Twilio secrets.",
    },
    {
        "name": "seed-platform-owner",
        "path": "supabase/functions/seed-platform-owner/index.ts",
        "endpoint": "POST /functions/v1/seed-platform-owner",
        "purpose": "Bootstrap utility: create/find auth user by email, create profile, assign platform_owner role, audit log, optional recovery link.",
        "auth": "No JWT verification observed — accepts JSON {email, redirectTo} with service role client only.",
        "roles": "Unrestricted caller if endpoint is public (CRITICAL).",
        "validation": "Email required in body.",
        "tables": "auth.users (admin), profiles, platform_roles, platform_audit_log.",
        "errors": "400 if email missing; throws on profile/role failures.",
        "security": "CRITICAL: must be network-restricted or removed in production; equivalent to open platform admin creation.",
        "production": "Not production-safe — disable or protect with service-role-only gateway.",
    },
]


def add_architecture_diagrams(doc):
    add_heading(doc, "1.1.1 System Architecture Diagram (ASCII)", 3)
    add_para(doc, "The following diagram summarizes runtime components and trust boundaries identified in the codebase:")
    add_mono_block(doc, ARCHITECTURE_DIAGRAM)
    add_heading(doc, "1.1.2 Authentication and Data Flow Diagram (ASCII)", 3)
    add_mono_block(doc, DATA_FLOW_DIAGRAM)


def add_er_diagram(doc):
    add_heading(doc, "4.2.1 Entity-Relationship Diagram (Tabular)", 3)
    add_para(doc, "Core relationships inferred from supabase/migrations and src/integrations/supabase/types.ts. Cardinality is logical, not every FK name listed.")
    add_table(doc, ["Parent Entity", "Cardinality", "Child / Related Entity", "Notes"], ER_RELATIONSHIPS)


def add_edge_function_detailed_audits(doc):
    add_heading(doc, "5.8 Detailed Per-Function Edge Function Audit (All 13)", 2)
    add_para(doc, "Supabase config.toml sets verify_jwt = false for all functions below; effective security depends on custom auth in each handler. Production-ready column reflects code review only, not live penetration testing.")
    for i, fn in enumerate(EDGE_FUNCTION_AUDITS, 1):
        add_heading(doc, f"5.8.{i} {fn['name']}", 3)
        add_detail_table(doc, [
            ["Source file", fn["path"]],
            ["HTTP endpoint", fn["endpoint"]],
            ["Purpose", fn["purpose"]],
            ["Authentication required?", fn["auth"]],
            ["Role restrictions", fn["roles"]],
            ["Input validation", fn["validation"]],
            ["Database tables touched", fn["tables"]],
            ["Error handling", fn["errors"]],
            ["Security concerns", fn["security"]],
            ["Production-ready?", fn["production"]],
        ])
    add_para(doc, "Summary: 2 functions flagged Not production-safe (process-reminders, seed-platform-owner); majority Partial pending secrets, RLS verification, and operational hardening.", bold=True)


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title page
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BlueThread Solutions")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(41, 82, 144)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Technical Platform Audit Report")
    r2.bold = True
    r2.font.size = Pt(20)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Interpreter Scheduling / Agency Management Platform",
        "",
        "Document type: Static codebase audit (read-only)",
        "Repository: bluethreadsolutions-main",
        "Audit date: May 23, 2026",
        "",
        "Methodology: Source inspection of frontend, Supabase migrations,",
        "Edge Functions, CI configuration, and test suites.",
        "No live Supabase project, deployment, or penetration test was",
        "performed unless explicitly stated.",
    ]:
        m = meta.add_run(line + "\n")
        m.font.size = Pt(11)
        m.italic = True

    doc.add_page_break()

    # Table of contents placeholder
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "Part A — Internal Developer Notes (Expanded)",
        "Part B — Client-Ready Technical Audit Report",
        "  1. Project Structure and Technology Stack",
        "  2. Installation, Build, and Local Run Readiness",
        "  3. Data Layer Architecture",
        "  4. Database and Data Model (Detailed)",
        "  5. Edge Functions — API Audit (Detailed)",
        "    5.8 Detailed Per-Function Audit (All 13 Functions)",
        "  1.1.1 System Architecture Diagram (ASCII)",
        "  1.1.2 Data Flow Diagram (ASCII)",
        "  4.2.1 Entity-Relationship Diagram (Tabular)",
        "  6. Authentication and Authorization (Detailed)",
        "  7. Scheduling and Calendar Logic (Detailed)",
        "  8. Payment and Billing (Detailed)",
        "  9. HIPAA-Conscious Security Review (Expanded)",
        "  10. General Security Audit (OWASP-Oriented)",
        "  11. Code Quality and Maintainability",
        "  12. Testing and QA Readiness",
        "  13. Deployment and Production Readiness",
        "  14. Frontend Feature Reality Check (Expanded)",
        "  15. Feature Completion Matrix (Expanded)",
        "  16. Security Risk Matrix (Full)",
        "  17. Recommended Development Plan",
        "  18. Reuse vs Rebuild (Detailed)",
        "  19. Open Questions for Client",
        "  20. Final Recommendation",
        "Part C — Appendices",
        "Part D — Verification Checklist",
    ]
    for item in toc_items:
        add_bullet(doc, item)
    doc.add_page_break()

    # ========== PART A ==========
    add_heading(doc, "Part A — Internal Developer Notes (Expanded)", 1)

    add_heading(doc, "A.1 Executive Technical Verdict", 2)
    add_table(doc, ["Claim", "Verdict"], [
        ["~80% UI built", "Plausible — ~40 page modules, ~104 component files, ~26 hooks"],
        ["Backend missing", "Incorrect — backend is Supabase-native (Postgres + RLS + 29 SQL functions + 13 Edge Functions)"],
        ["Production ready", "No — demo auth bypass, unauthenticated cron path, unverified live RLS, no MFA, payment model mismatch risk"],
        ["HIPAA compliant", "Cannot certify — Low technical readiness; patient_client_name and medical-style demo content present"],
    ])

    add_para(doc, "Architecture in one sentence: Single-page React app talks to Supabase PostgREST/RPC with privileged operations delegated to Deno Edge Functions using service-role clients after JWT validation in supabase/functions/_shared/cors.ts.", italic=True)

    add_heading(doc, "A.2 Repository Metrics (Confirmed from Filesystem)", 2)
    add_table(doc, ["Metric", "Value", "Path / Note"], [
        ["SQL migrations", "100", "supabase/migrations/"],
        ["RLS policies (CREATE POLICY)", "~133", "Across migrations"],
        ["Public SQL functions", "29", "See Section B.4.7"],
        ["Public tables (CREATE TABLE)", "45", "See Section B.4.3"],
        ["Edge Functions", "13 + _shared", "supabase/functions/"],
        ["Frontend TS/TSX files", "224", "src/"],
        ["Page modules", "40", "src/pages/"],
        ["Custom hooks", "26", "src/hooks/"],
        ["UI/components", "104", "src/components/"],
        ["useAdaptedQuery / mutation usages", "57", "Dual demo/production data path"],
        ["functions.invoke from frontend", "28 call sites", "See Section B.5.7"],
        ["Vitest test files", "1 (example.test.ts)", "~290 lines of logic tests"],
        ["Playwright specs", "11", "e2e/*.spec.ts"],
        ["Supabase generated types", "~4,261 lines", "src/integrations/supabase/types.ts"],
        ["node_modules in ZIP", "Absent", "Requires npm install"],
        [".env in ZIP", "Present", "Listed in directory; gitignored in .gitignore"],
    ])

    add_heading(doc, "A.3 Critical Path for Local Verification", 2)
    add_para(doc, "cd bluethreadsolutions-main")
    add_para(doc, "npm install")
    add_para(doc, "cp .env.example .env   # fill VITE_SUPABASE_* from Supabase dashboard")
    add_para(doc, "npm run dev            # Vite :8080 per vite.config.ts")
    add_para(doc, "npm run build")
    add_para(doc, "npm run lint")
    add_para(doc, "npm run test")
    add_para(doc, "# E2E (needs .env.test + live Supabase + seeded users):")
    add_para(doc, "cp .env.test.example .env.test")
    add_para(doc, "npm run test:e2e")
    add_para(doc, "Supabase (required for real data): supabase db push; supabase functions deploy; set secrets: RESEND_API_KEY, TWILIO_*, GOOGLE_CALENDAR_*, QBO_*, QBO_WEBHOOK_VERIFIER_TOKEN, APP_BASE_URL")

    add_heading(doc, "A.4 Top 10 Engineering Priorities", 2)
    priorities = [
        "Rotate secrets if ZIP contained .env with real keys.",
        "Disable /demo in production builds (import.meta.env.PROD gate).",
        "Remove unauthenticated branch in process-reminders (lines 147-148).",
        "RLS pen-test per role on staging (especially requester + interpreter).",
        "Server-side conflict constraint for interpreter double-booking.",
        "Clarify payments with client (QBO invoicing vs Stripe).",
        "Deploy Edge Function secrets (Resend, Twilio, Google, Intuit).",
        "MFA for agency_admin / scheduler.",
        "Expand automated tests beyond logic-only Vitest.",
        "Document migration order and rollback for 100 files.",
    ]
    for i, p in enumerate(priorities, 1):
        add_bullet(doc, f"{i}. {p}")

    doc.add_page_break()

    # ========== PART B ==========
    add_heading(doc, "Part B — Client-Ready Technical Audit Report", 1)

    add_heading(doc, "Executive Summary", 2)
    add_para(doc, "Overall status: The codebase is substantially more complete than a typical ~80% UI Lovable export. There is a real multi-tenant data model, Row Level Security, server-side Edge Functions, scheduling/billing/messaging/import/calendar integrations, and role-based portals. It is not production-ready as-is for a regulated or HIPAA-sensitive launch.")
    add_table(doc, ["Dimension", "Assessment"], [
        ["Architecture", "Frontend + Supabase BaaS (not a separate Express/Node API server)"],
        ["UI completeness", "High — 40+ routed pages in src/App.tsx"],
        ["Backend completeness", "Medium–High for core agency ops; depends on deployed Supabase + secrets"],
        ["Payment processing", "Invoicing + QuickBooks — not consumer card payments"],
        ["HIPAA-conscious readiness", "Low (technical safeguards incomplete; compliance cannot be certified from code alone)"],
        ["Production launch", "Blocked until environment, security hardening, QA against live backend, and operational controls"],
    ])
    add_para(doc, "Main risks: Demo mode bypassing security; Edge Functions with JWT verification disabled; unauthenticated cron-style endpoints; client-side-only scheduling conflict prevention; sensitive fields (patient_client_name) without evident PHI controls; missing MFA/session policies; .env in deliverable ZIP.", bold=True)
    add_para(doc, "Best path forward: Continue from this codebase if the client accepts Supabase + Lovable stack. Phase 1: deploy staging Supabase, run migrations, wire secrets, disable/limit demo mode. Phase 2: security/RLS audit + server-side scheduling rules. Phase 3: operational HIPAA/BAA decisions (hosting, vendors, policies) separate from code.")

    add_heading(doc, "Document Control", 2)
    add_table(doc, ["Field", "Value"], [
        ["Product name", "BlueThread Solution (Interpreter Management Platform)"],
        ["Codebase name", "bluethreadsolutions v0.0.0"],
        ["Generator footprint", "Lovable (lovable-tagger, @lovable.dev/cloud-auth-js, Lovable README)"],
        ["Primary domain (CORS)", "https://app.bluethreadsolution.com"],
        ["Supabase project id (config)", "jznpmbkmipajyhlivtgy in supabase/config.toml"],
    ])

    # Section 1
    add_heading(doc, "1. Project Structure and Technology Stack", 1)
    add_heading(doc, "1.1 High-Level Architecture", 2)
    add_para(doc, "The application is a browser-based SPA (React + Vite) that communicates with Supabase for authentication, database access (PostgREST and RPC), Realtime subscriptions, and Edge Functions. Privileged workflows (invitations, imports, reminders, QuickBooks, Google Calendar, platform administration) execute in Deno Edge Functions using service-role Supabase clients after JWT validation. Third-party integrations include Resend (email), Twilio (SMS), Google Calendar API, Intuit QuickBooks Online, and Lovable Cloud Auth for OAuth.")
    add_architecture_diagrams(doc)

    add_heading(doc, "1.2 Frontend Stack (Confirmed)", 2)
    add_table(doc, ["Component", "Version / Detail", "Evidence"], [
        ["Runtime", "React 18.3.1", "package.json"],
        ["Build", "Vite 5.4.19, @vitejs/plugin-react-swc", "vite.config.ts"],
        ["Language", "TypeScript 5.8 (strict: false)", "tsconfig.app.json"],
        ["Routing", "react-router-dom 6.30", "src/App.tsx"],
        ["Server state", "@tanstack/react-query 5.83", "App.tsx QueryClient"],
        ["Forms", "react-hook-form + zod", "dependencies"],
        ["UI", "shadcn/Radix, Tailwind 3.4, framer-motion", "components.json, pages"],
        ["DnD scheduling", "@dnd-kit/core", "Schedule Wizard"],
        ["Charts", "recharts", "Reports/Dashboard"],
        ["Excel import", "xlsx", "Import wizard"],
        ["Signatures", "react-signature-canvas", "Time/completion flows"],
        ["Path alias", "@/* -> ./src/*", "vite.config.ts, tsconfig"],
    ])
    add_para(doc, "Main entry: index.html -> src/main.tsx -> src/App.tsx wrapped in ErrorBoundary.")
    add_para(doc, "Dev server: host ::, port 8080 (vite.config.ts).")

    add_heading(doc, "1.3 Backend Model (Confirmed — Not a Separate Node API)", 2)
    add_para(doc, "There is no server/, api/, or Express/FastAPI application. Backend responsibilities are split across:")
    add_table(doc, ["Layer", "Responsibility"], [
        ["PostgreSQL", "Schema, constraints, triggers, RLS, SECURITY DEFINER RPCs"],
        ["Supabase Auth", "Users, sessions, password reset, OAuth bridge"],
        ["Edge Functions", "Privileged workflows (invite, import, reminders, QBO, GCal, platform admin)"],
        ["PostgREST", "CRUD from frontend with user JWT"],
    ])
    add_para(doc, "This is a standard Lovable/Supabase BaaS pattern, not an incomplete backend—it is a different architecture than a monolithic API server.")

    add_heading(doc, "1.4 Third-Party Integrations (Code-Confirmed)", 2)
    add_table(doc, ["Service", "Purpose", "Configuration"], [
        ["Supabase", "DB, auth, realtime, functions", "VITE_SUPABASE_URL, VITE_SUPABASE_PUBLISHABLE_KEY"],
        ["Resend", "Transactional email", "RESEND_API_KEY in Edge (delivery.ts)"],
        ["Twilio", "SMS", "TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN, TWILIO_PHONE_NUMBER"],
        ["Google Calendar", "OAuth + event sync", "GOOGLE_CALENDAR_CLIENT_ID/SECRET"],
        ["Intuit QuickBooks", "OAuth + invoice sync + webhooks", "QBO env vars in qbo-auth, qbo-webhook"],
        ["Lovable Cloud Auth", "Google (and Apple/Microsoft in API) OAuth", "@lovable.dev/cloud-auth-js"],
    ])
    add_para(doc, "Not found in codebase: Stripe, PayPal, Square, Outlook/Microsoft Graph calendar, Firebase, Prisma, Drizzle.", bold=True)

    add_heading(doc, "1.5 Environment Variables", 2)
    add_para(doc, "Frontend (documented in .env.example): VITE_SUPABASE_PROJECT_ID, VITE_SUPABASE_URL, VITE_SUPABASE_PUBLISHABLE_KEY")
    add_table(doc, ["Variable", "Used In"], [
        ["SUPABASE_URL", "All Edge Functions"],
        ["SUPABASE_ANON_KEY", "JWT validation"],
        ["SUPABASE_SERVICE_ROLE_KEY", "Privileged DB access"],
        ["RESEND_API_KEY", "_shared/delivery.ts"],
        ["TWILIO_*", "_shared/delivery.ts"],
        ["GOOGLE_CALENDAR_CLIENT_ID/SECRET", "google-calendar-sync"],
        ["QBO_WEBHOOK_VERIFIER_TOKEN", "qbo-webhook"],
        ["APP_BASE_URL", "invite-user (default https://app.bluethreadsolution.com)"],
    ])
    add_para(doc, "E2E: .env.test.example — role emails/passwords, BASE_URL.")
    add_para(doc, "Risk: .env file exists in delivered ZIP; .gitignore excludes it from git but not from ZIP exfiltration. Treat as potential secret exposure.", bold=True)

    add_heading(doc, "1.6 Routing Map (Authoritative: src/App.tsx)", 2)
    add_table(doc, ["Category", "Routes"], [
        ["Public", "/, /login, /signup, /reset-password, /demo, /join/:agencySlug"],
        ["Onboarding", "/onboarding, /pending-approval"],
        ["All authenticated roles", "/dashboard, /messages, /settings"],
        ["Admin + Scheduler", "/appointments, /schedule-wizard, /customers, /customers/:id, /interpreters"],
        ["Admin only", "/billing-rates, /customer-billing, /interpreter-pay, /invoices, /billing-report, /notification-templates, /notification-log, /audit-log, /reports, /calendar-settings, /regions, /import, /import-history, /qbo-sync-log, /integration-health"],
        ["Requester", "/request, /my-requests"],
        ["Interpreter", "/my-schedule, /my-earnings, /time-tracking, /my-languages, /availability, /available-jobs"],
        ["Platform owner", "/platform/dashboard, /platform/agencies, /platform/agencies/:id, /platform/users, /platform/revenue, /platform/support, /platform/feature-flags, /platform/diagnostics, /platform/audit, /platform/settings, /platform/runbook"],
    ])
    add_para(doc, "Role enforcement (client): ProtectedRoute + getRolesForPath() from src/lib/route-roles.ts.")
    add_para(doc, "Role enforcement (server): RLS policies + Edge Function authenticateCaller().")

    doc.add_page_break()

    # Section 2
    add_heading(doc, "2. Installation, Build, and Local Run Readiness", 1)
    add_heading(doc, "2.1 Package Management", 2)
    add_para(doc, "Primary: npm (package-lock.json present). Also present: bun.lock, bun.lockb (dual lockfiles — pick one package manager for CI to avoid drift).")

    add_heading(doc, "2.2 Scripts (package.json)", 2)
    add_table(doc, ["Script", "Command", "Purpose"], [
        ["dev", "vite", "Local dev :8080"],
        ["build", "vite build", "Production bundle"],
        ["build:dev", "vite build --mode development", "Dev-mode build"],
        ["lint", "eslint .", "ESLint 9"],
        ["preview", "vite preview", "Preview production build"],
        ["test", "vitest run", "Unit tests"],
        ["test:watch", "vitest", "Watch mode"],
        ["test:e2e", "playwright test", "E2E (3 browsers)"],
        ["test:e2e:ui", "playwright test --ui", "Interactive E2E"],
    ])

    add_heading(doc, "2.3 TypeScript Posture", 2)
    add_para(doc, "tsconfig.app.json: strict: false, noImplicitAny: false, unused locals/parameters not enforced. Impact: Faster AI/Lovable iteration; higher runtime risk. Recommend enabling strict incrementally before enterprise launch.")

    add_heading(doc, "2.4 Likely Install/Build Outcome", 2)
    add_table(doc, ["Step", "Expected Result", "Blockers"], [
        ["npm install", "Should succeed", "None observed in manifest"],
        ["npm run build", "Should succeed if env vars set", "Missing VITE_SUPABASE_* may yield undefined client URL"],
        ["npm run dev", "UI loads", "Data empty/errors without Supabase"],
        ["Functional app", "Requires live Supabase + migrations + function deploy", "No local Supabase seed in repo root"],
    ])

    add_heading(doc, "2.5 Missing Operational Documentation", 2)
    add_bullet(doc, "README is generic Lovable template (clone, npm i, npm run dev).")
    add_bullet(doc, "Missing: Supabase project linking, migration apply procedure, Edge Function secret matrix, staging vs production separation, platform owner bootstrap procedure.")

    # Section 3
    add_heading(doc, "3. Data Layer Architecture", 1)
    add_heading(doc, "3.1 Dual-Mode Data Access Pattern", 2)
    add_para(doc, "The codebase implements an explicit demo vs production strategy in src/lib/data-adapter.ts via useAdaptedQuery and useAdaptedMutation:")
    add_bullet(doc, "When isDemoMode is false: hooks call Supabase (queryFn / mutationFn).")
    add_bullet(doc, "When isDemoMode is true: hooks return in-memory data from DemoDataContext without network calls.")
    add_para(doc, "Hooks such as useCustomers, useInvoices, usePaginatedAppointments are production-capable when demo mode is off.")
    add_para(doc, "Sales demos can show full UX without database — must not be confused with production behavior.")

    add_heading(doc, "3.2 Demo Mode Mechanics", 2)
    add_table(doc, ["Mechanism", "Behavior", "File"], [
        ["Role selection", "sessionStorage.demo_role", "DemoContext.tsx"],
        ["Fake agency UUID", "demo-agency-00000000-0000-0000-0000-000000000000", "DemoContext.tsx"],
        ["Auth bypass", "ProtectedRoute returns children without Supabase user", "ProtectedRoute.tsx L26-31"],
        ["Fake session", "session: {} as Session", "AuthContext.tsx L104"],
    ])

    doc.add_page_break()

    # Section 4 Database
    add_heading(doc, "4. Database and Data Model (Detailed)", 1)
    add_heading(doc, "4.1 Migration Timeline", 2)
    add_para(doc, "100 migration files from 20260224* through 20260416* (~7 weeks of iterative schema evolution). Fresh deploy must apply all migrations in order; drift risk if Lovable cloud DB diverges from repo.")

    add_heading(doc, "4.2 Core Entity Model", 2)
    add_para(doc, "Multi-tenant model: agencies -> profiles, user_roles, customers -> locations, appointments (with interpreter/requester), billing_rates, invoices, conversations/messages, import batches, QBO connections, platform_* tables for SaaS operator console.")
    add_er_diagram(doc)

    add_heading(doc, "4.3 Tables Created in Migrations (45 public tables)", 2)
    tables_list = (
        "agencies, profiles, user_roles, languages, interpreter_languages, customers, locations, "
        "appointments, appointment_history, interpreter_availability, billing_rates, billing_rules, "
        "interpreter_pay_rates, interpreter_regions, interpreter_notes, interpreter_notification_prefs, "
        "invoices, invoice_line_items, conversations, conversation_participants, messages, notifications, "
        "notification_templates, notification_log, invitations, customer_requestors, requestor_locations, "
        "regions, google_calendar_connections, qbo_connections, qbo_item_mappings, qbo_sync_jobs, "
        "qbo_sync_log, qbo_webhook_events, import_batches, import_batch_rows, import_mapping_templates, "
        "import_mapping_rules, import_quality_thresholds, platform_roles, platform_audit_log, "
        "platform_billing_config, platform_invoices, platform_invoice_line_items, platform_qbo_connection, "
        "platform_qbo_sync_log, platform_usage_log, join_requests (+ views *_live referenced in types)."
    )
    add_para(doc, tables_list)

    add_heading(doc, "4.4 Appointment Domain Fields (Scheduling + PHI-Adjacent)", 2)
    add_table(doc, ["Field", "Purpose", "HIPAA Note"], [
        ["patient_client_name", "End client/patient name on appointment", "Potential identifier"],
        ["client_reference", "External reference ID", "May link to medical record"],
        ["description, notes, agency_notes, interpreter_notes", "Free text", "May contain clinical context"],
        ["custom_fields JSONB", "Extensibility + override logs", "May store sensitive metadata"],
        ["recurrence_rule, parent_recurring_id", "Series", "—"],
        ["gcal_event_id, gcal_sync_status", "Calendar sync", "Copies data to Google"],
        ["billing_breakdown, billed_amount", "Revenue", "—"],
        ["signature_data", "Completion capture", "Potential biometric-adjacent"],
        ["is_import_staged, source_*, is_deleted", "Data lifecycle", "—"],
    ])

    add_heading(doc, "4.5 Enums and Status Workflow", 2)
    add_para(doc, "Original enum (20260224164124): appointment_status: pending, scheduled, confirmed, in_progress, completed, cancelled, no_show.")
    add_para(doc, "Later migrations expanded statuses (process-import VALID_STATUSES): requested, requested_last_minute, interpreter_assigned, interpreter_assigned_last_minute, interpreter_confirmed, reassignment_needed, completed_last_minute, cancelled, late_cancel_no_show_client, no_show_interpreter, etc.")
    add_para(doc, "Assignment methods: self_claim, availability, offer, manual (+ admin_confirmed in wizard code).")

    add_heading(doc, "4.6 Row Level Security (RLS)", 2)
    add_para(doc, "RLS enabled on core tables from initial migration. ~133 policies across 100 files — pattern: agency_id = get_user_agency_id(auth.uid()) + has_role().")
    add_para(doc, "Documented security hardening (20260309220805): (1) Profiles: users cannot change own agency_id/customer_id on update. (2) appointment_history: removed direct INSERT policy that allowed forgery. (3) Requesters: INSERT requires requester_id = auth.uid() and customer_id matches get_user_customer_id().")
    add_para(doc, "Platform roles (20260324044254): is_platform_owner() SECURITY DEFINER; users can SELECT own platform role row.")
    add_para(doc, "Cannot verify without live DB: policy gaps, SECURITY DEFINER function abuse, view (appointments_live) bypass paths.")

    add_heading(doc, "4.7 Server-Side SQL Functions (29)", 2)
    add_table(doc, ["Function", "Likely Purpose"], [
        ["bootstrap_agency_admin", "Agency + admin provisioning on onboarding"],
        ["search_appointments", "Paginated, filtered appointment search"],
        ["get_report_data", "Admin reports aggregation (SECURITY DEFINER)"],
        ["get_dashboard_counts", "Dashboard metrics"],
        ["log_appointment_change", "Audit trigger helper"],
        ["has_role, get_user_agency_id, get_user_customer_id", "RLS helpers"],
        ["is_platform_owner, get_platform_*, search_platform_users", "Platform console"],
        ["check_import_concurrency, rollback_import_batch, transition_import_batch", "Import pipeline"],
        ["cancel_stale_reminders", "Reminder dedup/cleanup"],
        ["handle_new_user", "Profile creation on signup"],
        ["validate_billing_rate_customer, validate_pay_rate_interpreter", "Billing integrity"],
    ])
    add_para(doc, "search_appointments: Called from usePaginatedAppointments.ts via .rpc() with pagination, status filters, text search, assignment filter — real server-side query, not UI-only.")

    doc.add_page_break()

    # Section 5 Edge Functions
    add_heading(doc, "5. Edge Functions — API Audit (Detailed)", 1)
    add_heading(doc, "5.1 Global Configuration", 2)
    add_para(doc, "supabase/config.toml: All listed functions have verify_jwt = false. Supabase gateway does not enforce JWT; each function implements custom auth (or omits it).")

    add_heading(doc, "5.2 Shared Authentication (_shared/cors.ts)", 2)
    add_para(doc, "authenticateCaller(req): (1) Requires Authorization: Bearer <token>. (2) Validates JWT via anonClient.auth.getClaims(token). (3) Loads profiles.agency_id with service role. (4) Loads user_roles for that agency. (5) Returns { userId, agencyId, roles, adminClient }.")
    add_para(doc, "CORS allowlist: app.bluethreadsolution.com, specific Lovable preview hosts; regex https://*.lovable.app and https://*.lovableproject.com; unknown origins fall back to first allowed origin.")

    add_heading(doc, "5.3 Function Inventory", 2)
    add_table(doc, ["Function", "Auth Model", "Primary Actions"], [
        ["accept-invitation", "JWT + email match", "Link user to pending invite"],
        ["auth-email-hook", "Supabase hook", "Custom auth emails"],
        ["google-calendar-sync", "authenticateCaller", "OAuth, sync, delete events"],
        ["invite-user", "Admin/scheduler checks", "invite/create/resend/revoke"],
        ["manage-join-request", "Authenticated", "Agency join workflow"],
        ["platform-admin", "Platform owner / bootstrap / service seed", "Agency CRUD, promote owner, audit"],
        ["platform-qbo", "Platform owner", "Platform-level QBO"],
        ["process-import", "authenticateCaller", "dry_run, execute, rollback, resume"],
        ["process-reminders", "Service role OR admin JWT; none -> allow", "Email/SMS reminders"],
        ["qbo-auth", "Mixed", "OAuth start/callback/status"],
        ["qbo-webhook", "HMAC intuit-signature", "Entity change events"],
        ["send-notification", "Role-gated", "email/sms/in_app"],
        ["seed-platform-owner", "Service role", "Bootstrap platform user"],
    ])

    add_heading(doc, "5.4 Notification Delivery Pipeline", 2)
    add_para(doc, "_shared/delivery.ts: Email via Resend API with HTML escaping in buildBrandedHtml; SMS via Twilio REST; writes to notification_log. Reminder content (process-reminders) includes assignment title, datetime, location address, customer name, language — may be sensitive in SMS/email.")

    add_heading(doc, "5.5 Critical: process-reminders Authentication Gap", 2)
    add_para(doc, "Evidence: supabase/functions/process-reminders/index.ts lines 115-148. If no Authorization header, the call is allowed (backward-compatible cron invocations). Impact: Anyone who discovers the function URL can trigger reminder processing (email/SMS charges, appointment metadata reads via service role). Recommendation: Require Authorization: Bearer <service_role> or signed cron secret header; remove no-auth path before production.", bold=True)

    add_heading(doc, "5.6 QBO Webhook (Positive Pattern)", 2)
    add_para(doc, "qbo-webhook/index.ts: HMAC-SHA256 verification against QBO_WEBHOOK_VERIFIER_TOKEN; idempotency via qbo_webhook_events; realm -> agency mapping.")

    add_heading(doc, "5.7 Frontend -> Edge Function Invocation Map", 2)
    add_table(doc, ["Caller", "Function", "Typical Action"], [
        ["useAgencyData.ts", "send-notification, google-calendar-sync", "Appointment lifecycle"],
        ["useScheduleWizard.ts", "send-notification", "Assignment offers"],
        ["useImportWizard.ts, ImportHistory.tsx", "process-import", "CSV pipeline"],
        ["CalendarSettings.tsx, IntegrationHealth.tsx", "google-calendar-sync", "OAuth/sync test"],
        ["IntegrationHealth.tsx", "qbo-auth, send-notification", "Integration smoke tests"],
        ["Interpreters.tsx", "invite-user", "Staff onboarding"],
        ["Onboarding.tsx", "accept-invitation", "Post-signup linking"],
        ["JoinAgency.tsx", "manage-join-request", "Public join flow"],
        ["usePlatformData.ts", "platform-admin", "SaaS operator console"],
        ["MyRequests.tsx, MySchedule.tsx", "send-notification", "Requester/interpreter alerts"],
    ])

    doc.add_page_break()
    add_edge_function_detailed_audits(doc)
    doc.add_page_break()

    # Section 6 Auth
    add_heading(doc, "6. Authentication and Authorization (Detailed)", 1)
    add_table(doc, ["Mechanism", "Status", "Evidence"], [
        ["Email/password login", "Implemented", "Login.tsx -> signInWithPassword"],
        ["Sign up", "Implemented", "Signup.tsx -> signUp + email confirmation toast"],
        ["Password reset", "Implemented", "resetPasswordForEmail, RecoveryRedirect, AuthContext"],
        ["Google OAuth", "Partial", "GoogleSignInButton -> lovable.auth.signInWithOAuth"],
        ["Session storage", "localStorage via Supabase client", "client.ts"],
        ["MFA", "Not found", "—"],
        ["Email verification enforcement", "Not verified in code", "Supabase project setting dependent"],
        ["Account suspension", "profiles.is_active, agency agency_status", "Partial gating in ProtectedRoute"],
    ])

    add_heading(doc, "6.2 Role Model", 2)
    add_table(doc, ["Role", "app_role enum", "Route Access (Client)"], [
        ["Agency Admin", "agency_admin", "Full agency admin routes"],
        ["Scheduler", "scheduler", "Scheduling + customers + interpreters"],
        ["Requester", "requester", "/request, /my-requests"],
        ["Interpreter", "interpreter", "Schedule, availability, jobs, earnings, time tracking"],
        ["Platform Owner", "platform_roles.role = platform_owner", "/platform/* via PlatformGuard"],
    ])
    add_para(doc, "Roles stored in user_roles table (not JWT claims) — correct for server-side checks.")

    add_heading(doc, "6.3 Authorization Layers", 2)
    add_table(doc, ["Layer", "Enforced?", "Notes"], [
        ["React Router / ProtectedRoute", "Yes (UI)", "Demo bypass"],
        ["route-roles.ts", "UI only", "Same roles as sidebar"],
        ["Supabase RLS", "Yes (if policies correct)", "Must pen-test"],
        ["Edge Functions", "Yes (custom)", "Service role after JWT"],
        ["RPC SECURITY DEFINER", "Yes", "get_report_data etc. uses auth.uid() internally"],
    ])

    add_heading(doc, "6.4 Platform Owner Bootstrap Risk", 2)
    add_para(doc, "platform-admin supports bootstrap (first authenticated user can become platform owner if none exist) and seed_by_email (requires service role key). Recommendation: Run bootstrap once in controlled environment; disable or protect thereafter.")

    add_heading(doc, "6.5 IDOR / Cross-Tenant Considerations", 2)
    add_para(doc, "Positive patterns: accept-invitation ignores client-supplied agency_id; send-notification verifies targetProfile.agency_id === agencyId; hooks consistently .eq(agency_id, profile.agency_id). Residual risks: any RLS bug exposes entire agency tenant; service role functions bypass RLS; demo mode has no tenant isolation.")

    # Section 7 Scheduling
    add_heading(doc, "7. Scheduling and Calendar Logic (Detailed)", 1)
    add_table(doc, ["Capability", "Implementation", "Persistence", "Server Enforcement"], [
        ["Create/edit appointments", "useAppointmentMutations, forms", "appointments table", "RLS by role"],
        ["Paginated search", "search_appointments RPC", "Postgres", "SECURITY DEFINER scoped to agency"],
        ["Recurring series", "bulkCreate/bulkUpdate, parent_recurring_id", "Postgres", "No series-level lock observed"],
        ["Schedule Wizard dispatch", "DnD UI + useAssignAppointment", "Updates appointments", "Conflict check client-side"],
        ["Conflict detection", "Overlap query in useInterpreterSchedulesBatch", "—", "Admin can override with reason"],
        ["Override audit", "appointment_history action override_conflict", "Postgres", "Optional insert non-blocking"],
        ["Availability", "interpreter_availability + recurring logic", "Postgres", "Viewable by agency"],
        ["Timezone", "agency-timezone.ts helpers", "agencies.timezone", "UTC storage, local display"],
        ["Self-claim jobs", "AvailableJobs + statuses", "Postgres", "Status + RLS dependent"],
        ["Google Calendar", "Edge function + connection table", "google_calendar_connections", "OAuth secrets on server"],
        ["Outlook", "Not found", "—", "—"],
        ["ICS export", "ics-generator.ts", "Client-generated file", "Vitest covered"],
        ["Reminders", "process-reminders cron", "notification_log", "Dedup keys per appt/channel"],
    ])
    add_para(doc, "Conflict handling: Schedule Wizard loads overlapping appointments and availability. On assign, UI can require override reason; writes custom_fields.override_log and appointment_history — does not block at database level.")
    add_para(doc, "Appointment side effects (production): sendAppointmentNotifications -> send-notification; triggerCalendarSync -> google-calendar-sync; bulk recurring with parent_recurring_id; optimistic concurrency on series update via updated_at.")

    # Section 8 Billing
    add_heading(doc, "8. Payment and Billing (Detailed)", 1)
    add_para(doc, "Billing means agency-to-customer invoicing for interpretation services, not in-app card checkout.")
    add_table(doc, ["Component", "Description"], [
        ["billing_rates", "Complex rate cards (hourly, tiered, premiums, cancellation fees)"],
        ["billing-engine.ts", "Client-side calculation (~524 lines) with timezone-aware premiums"],
        ["invoices / invoice_line_items", "Generated invoices: draft, sent, paid, overdue, cancelled"],
        ["InterpreterPay", "Interpreter compensation rates"],
        ["QBO integration", "Sync invoices/customers to QuickBooks"],
    ])
    add_table(doc, ["Expected by Many Clients", "In Codebase?"], [
        ["Stripe / card checkout", "No"],
        ["PCI SAQ-A via hosted checkout", "No"],
        ["PayPal", "No"],
        ["ACH", "No"],
        ["Invoice PDF + QBO", "Yes (partial)"],
        ["Payment status on invoice", "Yes (invoices.status)"],
        ["Webhook reconciliation (QBO)", "Yes (HMAC verified)"],
    ])
    add_para(doc, "SampleBillingReport.tsx uses FAKE_AGENCY, FAKE_CUSTOMER (Atrium Health), and medical appointment titles — marketing/demo only. Must not ship as production reporting.")

    doc.add_page_break()

    # Section 9 HIPAA
    add_heading(doc, "9. HIPAA-Conscious Security Review (Expanded)", 1)
    add_para(doc, "DISCLAIMER: This audit does not certify HIPAA compliance. Compliance requires organizational policies, risk analysis, workforce training, BAAs, and operational controls beyond source code.", bold=True)

    add_heading(doc, "9.2 HHS Security Rule Mapping (Technical Lens)", 2)
    add_table(doc, ["Safeguard Area", "Code Evidence", "Gap"], [
        ["Access control", "RLS, roles, is_active", "No MFA; demo bypass; no fine-grained PHI role"],
        ["Audit controls", "appointment_history, platform_audit_log, notification_log", "No PHI access log; no log integrity guarantees"],
        ["Integrity", "updated_at optimistic locking on series", "No cryptographic integrity"],
        ["Person/entity authentication", "Supabase Auth", "No MFA; weak session policy"],
        ["Transmission security", "HTTPS assumed (Supabase/Vite)", "SMS/email content may include identifiers"],
        ["Encryption at rest", "Supabase platform responsibility", "Not verifiable in repo"],
        ["Minimum necessary", "Broad appointment fields", "patient_client_name, notes, reminders include location/customer"],
    ])

    add_heading(doc, "9.3 Data That May Constitute PHI (Client-Dependent)", 2)
    for item in [
        "patient_client_name",
        "Appointment title/description (e.g. Cardiology Follow-up in samples)",
        "Location addresses (hospital/clinic)",
        "Contact info in profiles/customers",
        "Messages between requester and interpreter",
        "Signature capture on completion",
        "Google Calendar event copies",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9.4 Vendor BAA Checklist (Must Verify Contractually)", 2)
    add_table(doc, ["Vendor", "Data Touched"], [
        ["Supabase", "All DB"],
        ["Resend", "Email content"],
        ["Twilio", "SMS content"],
        ["Google", "Calendar events"],
        ["Intuit", "Invoice/customer financial"],
        ["Lovable / hosting", "App delivery, possibly OAuth"],
    ])

    add_para(doc, "HIPAA-Conscious Readiness Score: LOW for covered-entity deployment without additional controls.", bold=True)

    # Section 10 Security
    add_heading(doc, "10. General Security Audit (OWASP-Oriented)", 1)
    add_heading(doc, "10.1 Critical", 2)
    add_table(doc, ["ID", "Finding", "Evidence", "Recommendation", "Priority"], [
        ["SEC-01", "Demo auth bypass", "ProtectedRoute.tsx", "Disable /demo in prod", "P0"],
        ["SEC-02", "Unauthenticated reminders", "process-reminders/index.ts:147", "Auth required", "P0"],
        ["SEC-03", "RLS not verified live", "100 migrations", "Pen-test", "P0"],
        ["SEC-04", "Leaked .env in ZIP", "Filesystem", "Rotate secrets", "P0"],
    ])
    add_heading(doc, "10.2 High", 2)
    add_table(doc, ["ID", "Finding", "Evidence", "Recommendation", "Priority"], [
        ["SEC-05", "JWT verification off on all Edge Functions", "config.toml", "Per-function review", "P1"],
        ["SEC-06", "UI-only RBAC", "route-roles.ts", "Treat RLS as authoritative", "P1"],
        ["SEC-07", "Double booking", "useScheduleWizard.ts", "DB constraint", "P1"],
        ["SEC-08", "No MFA", "—", "Enable Supabase MFA", "P1"],
        ["SEC-09", "Broad Lovable CORS regex", "cors.ts", "Restrict to prod domain", "P1"],
        ["SEC-10", "Platform bootstrap", "platform-admin", "One-time controlled seed", "P1"],
    ])
    add_heading(doc, "10.3 Medium", 2)
    add_table(doc, ["ID", "Finding", "Evidence", "Recommendation", "Priority"], [
        ["SEC-11", "strict: false TypeScript", "tsconfig.app.json", "Enable strict mode", "P2"],
        ["SEC-12", "Message realtime + soft delete", "useMessages.ts", "XSS review on render", "P2"],
        ["SEC-13", "Service role in Edge Functions", "All privileged functions", "Minimize scopes, audit logs", "P2"],
        ["SEC-14", "Reminder SMS content", "buildReminderMessage", "Minimize content", "P2"],
        ["SEC-15", "Dependency audit not run", "—", "npm audit in CI", "P2"],
    ])
    add_heading(doc, "10.4 Low", 2)
    add_table(doc, ["ID", "Finding", "Evidence", "Recommendation", "Priority"], [
        ["SEC-16", "Dual lockfiles npm/bun", "root", "Standardize", "P3"],
        ["SEC-17", "Playwright only on failure artifacts", "CI workflow", "Add pass artifacts optional", "P3"],
    ])

    doc.add_page_break()

    # Section 11-13
    add_heading(doc, "11. Code Quality and Maintainability", 1)
    add_heading(doc, "11.1 Strengths", 2)
    for s in [
        "Centralized route-role map shared by router and sidebar",
        "data-adapter pattern reduces demo/production branching duplication",
        "Rich domain modeling in SQL (import staging, soft delete, audit history)",
        "Edge Function shared modules (handler.ts, cors.ts, delivery.ts)",
        "Timezone discipline documented in agency-timezone.ts",
        "Iterative security migrations (requester INSERT tightening, profile UPDATE fix)",
    ]:
        add_bullet(doc, s)
    add_heading(doc, "11.2 Technical Debt", 2)
    add_table(doc, ["Area", "Issue"], [
        ["Type safety", "strict: false, (supabase as any) in places"],
        ["File size", "Large pages/hooks (useAgencyData.ts 677+ lines, ScheduleWizard.tsx 660+ lines)"],
        ["Demo vs prod", "Two parallel universes — easy to mis-test"],
        ["Error handling", "Some console.warn non-blocking audit failures"],
        ["Documentation", "README not project-specific"],
        ["Generated types", "Must regenerate after migrations"],
    ])

    add_heading(doc, "12. Testing and QA Readiness", 1)
    add_table(doc, ["Suite", "Files", "Coverage Focus"], [
        ["Vitest", "src/test/example.test.ts", "ICS generation, message unread logic, dedup, notification recipients, appointment thread participants, conversation filter"],
        ["Playwright", "11 specs in e2e/", "auth, appointments, billing, customers, demo, interpreter, messaging, requester, responsive, settings, admin-nav"],
    ])
    add_para(doc, "CI (.github/workflows/e2e-tests.yml): Trigger push/PR to main; Node 20, npm ci, Playwright; secrets per-role emails/passwords + BASE_URL. Does not run npm run test (Vitest) or npm run lint in workflow observed.")
    add_heading(doc, "12.3 QA Gaps", 2)
    for g in [
        "No RLS automated tests",
        "No Edge Function integration tests",
        "No QBO webhook tests",
        "No load tests for import or search RPC",
        "No security regression suite",
        "E2E depends on external staging credentials",
    ]:
        add_bullet(doc, g)
    add_heading(doc, "12.4 Recommended Pre-Launch QA Plan", 2)
    qa_plan = [
        "Role matrix manual test (50+ cases): CRUD per entity per role",
        "Cross-tenant negative tests (two agencies)",
        "Scheduling: double-book attempt via API",
        "Import: 10k row CSV dry-run/execute/rollback",
        "Calendar: OAuth + create/update/cancel sync",
        "QBO: sandbox connect + webhook replay",
        "Notification: email/SMS opt-out paths",
        "Accessibility smoke (Radix helps; still verify)",
        "Mobile responsive (Schedule Wizard <768px per comments)",
        "Disaster: migration rollback drill on staging clone",
    ]
    for i, q in enumerate(qa_plan, 1):
        add_bullet(doc, f"{i}. {q}")

    add_heading(doc, "13. Deployment and Production Readiness", 1)
    add_table(doc, ["Tier", "Suggested Hosting"], [
        ["Frontend static", "Lovable publish, Vercel, Netlify, or Cloudflare Pages"],
        ["API/DB/Auth", "Supabase project (Pro for production features)"],
        ["Edge Functions", "Supabase Edge (same project)"],
        ["Cron", "Supabase pg_cron + pg_net (extension enabled in 20260311043210 migration)"],
    ])
    add_heading(doc, "13.2 Production Launch Blockers", 2)
    blockers = [
        "Supabase production project with all migrations applied",
        "All Edge secrets configured",
        "Custom domain + HTTPS on frontend",
        "CORS updated to production origin only",
        "Demo mode disabled",
        "process-reminders auth fixed",
        "RLS pen-test sign-off",
        "Client sign-off on QBO vs card payments",
        "HIPAA/legal sign-off if applicable",
        "Monitoring (Sentry/Logflare/etc.) — not in repo",
    ]
    for b in blockers:
        add_bullet(doc, b)

    doc.add_page_break()

    # Section 14 Frontend table - full
    add_heading(doc, "14. Frontend Feature Reality Check (Expanded)", 1)
    fe_rows = [
        ["Landing /", "UI-only", "Landing.tsx", "Marketing claims vs QBO reality", "Align marketing"],
        ["Login", "Functional", "signInWithPassword", "Needs Supabase", "Configure auth"],
        ["Signup", "Functional", "signUp + 8-char password", "Open registration policy", "Confirm invite-only"],
        ["Reset password", "Functional", "resetPasswordForEmail", "Email deliverability", "Configure SMTP/Resend hook"],
        ["Demo /demo", "UI-only", "DemoContext, adapter bypass", "Full auth bypass", "Disable prod"],
        ["Onboarding", "Partial/Functional", "bootstrap_agency_admin RPC, accept-invitation", "Enterprise plan unclear", "E2E new agency"],
        ["Dashboard", "Partial", "get_dashboard_counts RPC + demo", "—", "Verify widgets"],
        ["Appointments", "Partial", "search_appointments, mutations", "Large surface", "QA all statuses"],
        ["Schedule Wizard", "Partial", "Real mutations + notifications", "Client conflicts", "Server rules"],
        ["Customers/Detail", "Partial", "useCustomers, locations, requestors", "—", "RLS test"],
        ["Interpreters", "Partial", "invite-user invoke", "—", "Invite E2E"],
        ["Request interpreter", "Partial", "useAppointmentMutations", "—", "Requester E2E"],
        ["My Requests/Schedule", "Partial", "Role-filtered queries", "—", "E2E"],
        ["Availability", "Partial", "interpreter_availability + demo", "Recurring DST", "TZ tests"],
        ["Available Jobs", "Partial", "Self-claim statuses", "—", "Interpreter E2E"],
        ["Time tracking", "Unknown", "Page exists", "Not fully traced", "Dedicated review"],
        ["Billing rates", "Partial", "useBillingRates", "Complex engine", "Unit test engine"],
        ["Invoices", "Partial", "useInvoices, generate", "No card pay", "Client expectation"],
        ["Billing report", "Partial", "get_report_data RPC", "—", "Validate numbers"],
        ["Sample billing report", "UI-only FAKE", "SampleBillingReport.tsx", "Fake hospital data", "Remove prod"],
        ["Messages", "Partial", "Supabase + Realtime; demo static", "XSS/policy", "Security review"],
        ["Notifications admin", "Partial", "Templates + Edge send", "Provider keys", "Configure"],
        ["Calendar settings", "Partial", "google-calendar-sync", "OAuth setup", "Google Cloud project"],
        ["Audit log", "Partial", "appointment_history", "Not access log", "Extend if HIPAA"],
        ["Import wizard", "Partial", "process-import ~1200+ lines", "Concurrency", "Load test"],
        ["QBO sync log", "Partial", "Tables + UI", "Sandbox", "Intuit setup"],
        ["Integration health", "Partial", "Smoke invokes", "Admin only", "Staging"],
        ["Reports", "Partial", "get_report_data", "—", "Cross-check"],
        ["Regions", "Partial", "regions table", "Feature flag", "—"],
        ["Settings", "Partial", "Agency settings hooks", "—", "QA"],
        ["Platform console", "Partial", "platform-admin", "Bootstrap risk", "Lock down"],
        ["Join agency", "Partial", "manage-join-request", "Public route", "Abuse review"],
        ["My earnings", "Partial", "Interpreter billing views", "—", "QA"],
        ["Interpreter pay", "Partial", "Admin rate management", "—", "QA"],
        ["Customer billing assignment", "Partial", "Admin UI", "—", "QA"],
        ["Import history", "Partial", "Rollback invoke", "—", "QA"],
        ["Pending approval", "Partial", "agency_status gating", "—", "QA"],
        ["Not found", "UI", "NotFound.tsx", "—", "—"],
    ]
    add_table(doc, ["Feature / Screen", "Status", "Evidence", "Main Issue", "Action"], fe_rows)

    doc.add_page_break()

    # Section 15 Feature matrix
    add_heading(doc, "15. Feature Completion Matrix (Expanded)", 1)
    fm_rows = [
        ["Authentication", "Partial", "Supabase Auth + contexts", "Demo bypass; no MFA", "Harden + MFA", "M"],
        ["Registration", "Partial", "Signup.tsx", "Open signup?", "Policy + hooks", "L"],
        ["Password reset", "Partial", "Full flow", "Email config", "Test", "L"],
        ["Role management", "Partial", "user_roles, invites", "RLS unverified", "Pen-test", "M"],
        ["Admin dashboard", "Partial", "Dashboard + RPCs", "—", "QA", "M"],
        ["Interpreter dashboard", "Partial", "Multiple pages", "—", "E2E", "M"],
        ["Requester dashboard", "Partial", "Request/my-requests", "Data isolation", "RLS test", "M"],
        ["Scheduling core", "Partial", "Appointments + wizard", "Double-book", "Server constraints", "H"],
        ["Availability", "Partial", "Table + UI", "DST edge cases", "Test matrix", "M"],
        ["Calendar sync (Google)", "Partial", "Edge function", "Secrets", "OAuth project", "H"],
        ["Calendar sync (Outlook)", "Missing", "Not found", "Expectation gap", "New integration", "H"],
        ["Assignment engine", "Partial", "Wizard + statuses", "Client-only conflicts", "RPC", "M"],
        ["Notifications", "Partial", "Resend/Twilio/EF", "Unauth cron", "Fix reminders", "M"],
        ["Payments (card)", "Missing", "No Stripe", "Expectation", "New module", "H"],
        ["Invoicing", "Partial", "Tables + UI + engine", "Calc accuracy", "Test + QBO", "M"],
        ["QBO integration", "Partial", "auth/webhook/sync", "Sandbox->prod", "Intuit app review", "H"],
        ["Webhooks (QBO)", "Partial", "HMAC verify", "Replay tests", "QA", "M"],
        ["DB schema", "Partial", "100 migrations", "Drift risk", "Apply + diff", "M"],
        ["RLS", "Partial", "133 policies", "Unverified", "Audit", "H"],
        ["Audit logs", "Partial", "appointment_history", "Incomplete", "Extend", "M"],
        ["Import/ETL", "Partial", "process-import", "Data corruption", "Staged rollout", "H"],
        ["Platform SaaS layer", "Partial", "Platform routes", "Owner bootstrap", "Ops procedure", "M"],
        ["Messaging", "Partial", "Realtime", "Content policy", "Moderation", "M"],
        ["HIPAA safeguards", "Low", "PHI-like fields", "Regulatory", "Program, not patch", "H"],
        ["Deployment", "Partial", "Lovable + CI E2E", "No prod monitoring", "Full pipeline", "M"],
        ["Testing", "Low", "1 unit + 11 e2e", "Gaps", "Expand CI", "M"],
        ["Documentation", "Missing", "Generic README", "Onboarding cost", "Write runbooks", "M"],
    ]
    add_table(doc, ["Feature", "Status", "Evidence", "Risk", "Work Required", "Complexity"], fm_rows)

    # Section 16 Security matrix full
    add_heading(doc, "16. Security Risk Matrix (Full)", 1)
    sr_rows = [
        ["R1", "Demo mode auth bypass", "Critical", "High if /demo public", "ProtectedRoute.tsx", "Build-time disable", "P0"],
        ["R2", "Open reminder endpoint", "Critical", "Medium", "process-reminders", "Auth required", "P0"],
        ["R3", "RLS misconfiguration", "Critical", "Unknown", "Migrations only", "Pen-test", "P0"],
        ["R4", "Leaked .env in ZIP", "Critical", "If shared", "Filesystem", "Rotate secrets", "P0"],
        ["R5", "No MFA", "High", "High", "—", "Enable MFA", "P1"],
        ["R6", "Double booking", "High", "Medium", "Wizard client logic", "DB constraint", "P1"],
        ["R7", "Edge JWT off", "High", "Low-Med", "config.toml", "Review each", "P1"],
        ["R8", "PHI in SMS/email", "Medium", "High", "buildReminderMessage", "Minimize content", "P2"],
        ["R9", "Google calendar data residency", "Medium", "Med", "GCal sync", "BAA + DLP", "P2"],
        ["R10", "strict: false bugs", "Medium", "Med", "tsconfig", "Enable strict", "P2"],
        ["R11", "Platform owner bootstrap", "Medium", "Low", "platform-admin", "Disable post-seed", "P2"],
        ["R12", "Supply chain", "Medium", "Med", "npm deps", "Audit + pin", "P2"],
        ["R13", "XSS in messages", "Medium", "Low-Med", "Messages.tsx", "Sanitize render", "P2"],
        ["R14", "Import CSV injection", "Medium", "Low", "process-import", "Validate rows", "P2"],
        ["R15", "CORS wildcard Lovable", "Low-Med", "Med", "cors.ts", "Tighten prod", "P3"],
    ]
    add_table(doc, ["ID", "Risk", "Severity", "Likelihood", "Evidence", "Mitigation", "Priority"], sr_rows)

    doc.add_page_break()

    # Section 17-20
    add_heading(doc, "17. Recommended Development Plan (Phased, with Estimates)", 1)
    add_table(doc, ["Phase", "Duration (Indicative)", "Deliverables"], [
        ["1. Stabilize", "1-2 weeks", "Staging Supabase, migrations, secrets, local runbook, smoke tests"],
        ["2. Demo isolation", "3-5 days", "Prod build excludes demo; sales environment only"],
        ["3. Security hardening", "2-3 weeks", "RLS audit, reminders auth, MFA, session timeout, CORS"],
        ["4. Scheduling integrity", "2 weeks", "Server-side conflict RPC, recurrence QA, DST tests"],
        ["5. Integrations", "3-4 weeks", "GCal production OAuth, QBO sandbox->prod, monitoring"],
        ["6. Payments clarity", "0-6 weeks", "QBO completion OR Stripe greenfield"],
        ["7. Compliance program", "Parallel", "BAAs, policies, PHI logging, retention (if HIPAA)"],
        ["8. QA + launch", "2-3 weeks", "Expanded CI, load test, UAT, go-live"],
    ])
    add_para(doc, "Total indicative range: 10-18 weeks to production-hardened launch (team-dependent), excluding full HIPAA certification program.")

    add_heading(doc, "18. Reuse vs Rebuild (Detailed)", 1)
    add_heading(doc, "Keep (High Value)", 2)
    for k in [
        "Entire React UI shell and design system",
        "Supabase schema as domain baseline (100 migrations)",
        "data-adapter + React Query patterns",
        "Edge Function inventory (modify, don't discard)",
        "Billing calculation engine (add tests, don't rewrite blindly)",
        "Schedule Wizard UX and hooks structure",
        "Import pipeline (with hardening)",
    ]:
        add_bullet(doc, k)
    add_heading(doc, "Refactor", 2)
    for r in [
        "Demo system -> explicit VITE_ENABLE_DEMO=true",
        "useAgencyData.ts -> split by domain module",
        "TypeScript strict mode adoption",
        "Edge Function auth -> unified middleware; remove no-auth paths",
        "Consolidate notification triggers (reduce duplicate invoke sites)",
    ]:
        add_bullet(doc, r)
    add_heading(doc, "Rebuild or Add Greenfield", 2)
    for b in [
        "Card payments (Stripe) if in scope",
        "Outlook calendar if required",
        "HIPAA access logging product feature",
        "Enterprise SSO/SAML if required (not present)",
        "Production observability stack",
    ]:
        add_bullet(doc, b)
    add_heading(doc, "Do Not Rebuild", 2)
    add_bullet(doc, "Core Postgres schema from scratch unless fundamental design flaw found (none observed at architecture level).")
    add_bullet(doc, "UI from scratch (would discard 6+ weeks of Lovable output).")

    add_heading(doc, "19. Open Questions for Client (Scope / Quote)", 1)
    questions = [
        "Regulatory scope: HIPAA covered entity, business associate, or general PII only?",
        "Payments: Agency invoicing + QBO sufficient, or patient/client card payments at booking?",
        "Hosting mandate: Must stay on Supabase/Lovable or migrate to AWS?",
        "Calendar: Google only, or Microsoft 365/Outlook required day one?",
        "Demo mode: Public on marketing site or internal sales only?",
        "Identity: MFA, SSO (Okta/Azure AD), or password-only acceptable?",
        "Existing production: Is jznpmbkmipajyhlivtgy the production project? Any live customer data?",
        "SMS/email: Which vendors are contracted (Twilio/Resend BAAs)?",
        "Interpreter credentials: Document upload / certification tracking required?",
        "SLA & support: 24/7 on-call for medical scheduling failures?",
        "Data retention: How long must appointments/messages be kept?",
        "Multi-region: US-only data residency required?",
    ]
    for i, q in enumerate(questions, 1):
        add_bullet(doc, f"{i}. {q}")

    add_heading(doc, "G. Priority Fix List (Grouped)", 1)
    add_heading(doc, "Critical Launch Blockers", 3)
    for item in [
        "Deploy Supabase + migrations + Edge Function secrets",
        "Rotate any credentials if .env was shared in ZIP",
        "Disable or restrict demo mode in production",
        "Fix unauthenticated process-reminders invocation",
        "RLS verification (professional pen-test)",
        "Clarify payment model with client (QBO invoicing vs card payments)",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "High Priority", 3)
    for item in [
        "MFA for privileged roles",
        "Server-side double-booking prevention",
        "Complete Google Calendar + QBO configuration docs",
        "Edge Function auth review (verify_jwt = false everywhere)",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Medium Priority", 3)
    for item in [
        "Session idle timeout",
        "Expanded integration/E2E tests",
        "Production monitoring/error reporting (Sentry, etc. — not found in codebase)",
        "Remove hardcoded marketing medical examples from client-facing samples",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Nice-to-Have", 3)
    for item in ["Outlook calendar", "Advanced analytics", "Mobile PWA polish"]:
        add_bullet(doc, item)

    add_heading(doc, "C. What Is Already Built (Summary)", 1)
    add_heading(doc, "Confirmed Usable (When Supabase Deployed)", 3)
    for item in [
        "Multi-tenant agencies with profiles and roles (agency_admin, scheduler, requester, interpreter)",
        "Customers, locations, interpreters, languages, regions",
        "Appointments with status workflow, soft delete, import staging, history/audit",
        "Schedule Wizard dispatch board with assignment and conflict override logging",
        "Interpreter availability (recurring + specific dates)",
        "Requester portal for creating/updating own requests (RLS tightened in migrations)",
        "Invoicing & billing rates with calculation engine (src/lib/billing-engine.ts)",
        "QBO connect/sync/webhooks (signature verification in qbo-webhook)",
        "Google Calendar two-way sync infrastructure",
        "Messaging (conversations, messages; useMessages.ts hits Supabase + Realtime)",
        "Notifications templates, log, reminders cron",
        "CSV import wizard with dry-run/execute/rollback Edge Function",
        "Platform owner console (agencies, users, revenue, feature flags, audit)",
        "User invitations (invite-user, accept-invitation)",
        "Audit log UI for appointment changes",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Partial", 3)
    for item in [
        "Google Sign-In — UI exists; depends on Lovable Cloud Auth configuration",
        "Calendar sync — backend function exists; requires GOOGLE_CALENDAR_CLIENT_ID/SECRET",
        "QBO — substantial code; needs Intuit app credentials and webhooks",
        "E2E tests — exist but require deployed app + role credentials",
        "Recurring appointments — schema/UI components; verify end-to-end against live DB",
        "Reports/analytics — pages present; depth of real aggregation not fully verified per screen",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "UI-Only / Demo", 3)
    for item in [
        "/demo mode — in-memory data, no Supabase (DemoContext, DemoDataContext, data-adapter)",
        "SampleBillingReport.tsx — hardcoded FAKE_* medical-style line items",
        "Messages in demo — static placeholder (Messages.tsx line 174)",
        "Landing page — marketing content; not operational backend",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "D. Major Technical Gaps (Summary)", 1)
    add_table(doc, ["Area", "Gap"], [
        ["Backend", "No standalone API server; all logic is Supabase RLS + RPC + Edge Functions — must be deployed and configured"],
        ["Database", "Migrations exist locally; live project state unknown from ZIP alone"],
        ["Auth", "No MFA; no idle session timeout; demo bypass; platform bootstrap for first platform owner"],
        ["Security", "Edge Functions verify_jwt = false; process-reminders allows no-auth calls; CORS allows *.lovable.app"],
        ["Scheduling", "Conflict checks are client-side; overrides logged but double-booking still possible via direct API"],
        ["Payments", "No card payment / Stripe — only invoicing + QBO; payment processing may be misaligned with client expectations"],
        ["Calendar", "Outlook not found in codebase"],
        ["HIPAA", "No technical controls package (MFA, access logging for PHI views, encryption attestations, retention jobs, BAA tracking)"],
        ["Deployment", "No production runbook in repo beyond platform runbook page; secrets management not documented"],
        ["Documentation", "README is Lovable generic; missing Supabase migration/deploy guide"],
    ])

    add_heading(doc, "E. Security & HIPAA-Conscious Review (Summary)", 1)
    add_para(doc, "Readiness: Low for environments handling identifiable patient/client names or health-adjacent appointment context.")
    add_heading(doc, "Critical Concerns (Code-Evidenced)", 3)
    add_table(doc, ["Finding", "Evidence", "Impact", "Recommendation", "Priority"], [
        ["Demo mode bypasses auth", "ProtectedRoute.tsx lines 26-31; DemoContext.tsx", "Anyone can use app without login in demo", "Disable in production builds or env-gate /demo", "Critical"],
        ["Unauthenticated reminder runner", "process-reminders/index.ts lines 147-148", "Abuse could trigger mass email/SMS", "Require service role or signed cron secret only", "Critical"],
        ["Edge Functions JWT off", "supabase/config.toml", "Broader attack surface if custom auth fails", "Review each function; enable JWT where possible", "High"],
        ["Route protection is UI-only", "route-roles.ts, ProtectedRoute.tsx", "Direct API access if RLS weak", "Pen-test all tables as each role", "Critical"],
        [".env in ZIP", "File listing (gitignored but present)", "Credential leak if ZIP shared", "Rotate keys; never redistribute", "Critical"],
        ["No MFA", "Not found in codebase", "Weak authentication for sensitive data", "Enable Supabase MFA / SSO policy", "High"],
        ["No session timeout", "Not found", "Unattended session risk", "Implement idle logout + Supabase session limits", "High"],
        ["Client-side conflict override", "useScheduleWizard.ts", "Double-booking, billing errors", "DB constraints or RPC with locking", "High"],
        ["PHI in notifications/logs", "process-reminders builds body with location/customer", "Disclosure via email/SMS", "Minimize content; configurable redaction", "Medium"],
    ])
    add_para(doc, "What cannot be confirmed from code alone: Supabase project tier & BAA, email provider BAA, SMS provider BAA, Google/Intuit BAAs, hosting SSL, backup/DR, data retention legal requirements, whether data is PHI under client's legal analysis.")

    add_heading(doc, "F. Recommended Development Plan (Phases)", 1)
    phases = [
        ("Phase 1 — Stabilize and run locally (1–2 weeks)", "Provision Supabase project; run all migrations; copy .env.example; deploy Edge Functions; npm run dev smoke-test login and one appointment CRUD per role."),
        ("Phase 2 — Clean architecture and remove mock/demo (1 week)", "Env-flag demo routes off in production; remove or isolate SampleBillingReport; document demo vs production for sales."),
        ("Phase 3 — Build/complete backend and database (2–4 weeks)", "RLS penetration test per role; server-side appointment conflict RPC; lock down process-reminders and process-import."),
        ("Phase 4 — Secure authentication and roles (1–2 weeks)", "MFA for admin/scheduler; session idle timeout; email verification enforcement; account suspension flows."),
        ("Phase 5 — Complete scheduling and calendar (2–3 weeks)", "End-to-end Google Calendar OAuth in staging; recurring series QA; timezone/DST test matrix."),
        ("Phase 6 — Complete payment integration (2–4 weeks)", "If card payments needed: design Stripe Checkout + webhooks (new work). If QBO-only: complete connect, sync, webhook reconciliation."),
        ("Phase 7 — Audit logs and security safeguards (ongoing)", "PHI access logging, admin audit export, data retention/deletion jobs, minimum-necessary notifications, security headers, CSP, dependency audit."),
        ("Phase 8 — QA/testing (2–3 weeks)", "Expand automated tests; Playwright against staging; load test scheduling peaks; staging to production promotion checklist."),
        ("Phase 9 — Deployment/staging/production launch (2–3 weeks)", "Separate Supabase projects; custom domain; monitoring; production secrets management; rollback strategy."),
    ]
    for title, desc in phases:
        add_heading(doc, title, 3)
        add_para(doc, desc)

    add_heading(doc, "20. Final Recommendation", 1)
    add_para(doc, "Continue from this codebase as the foundation for an interpreter agency scheduling SaaS. The project is materially beyond a UI prototype but is not launch-ready for regulated healthcare or high-assurance environments without a dedicated hardening phase.")
    add_table(doc, ["Scenario", "Recommendation"], [
        ["Sales demos only", "Enable demo mode; do not connect production PHI"],
        ["Commercial launch (non-HIPAA)", "2-3 month harden + QBO/GCal + QA"],
        ["HIPAA-sensitive launch", "Add compliance program + 3-6+ months technical controls"],
        ["Needs Stripe at launch", "Plan new payment module; do not assume existing code"],
    ])

    doc.add_page_break()

    # Part C Appendix
    add_heading(doc, "Part C — Appendix: Route-to-Role Reference", 1)
    add_para(doc, "From src/lib/route-roles.ts (authoritative for UI):")
    add_table(doc, ["Path", "Roles"], [
        ["/onboarding", "Any authenticated"],
        ["/dashboard, /messages, /settings", "All four agency roles"],
        ["/appointments, /schedule-wizard, /customers, /customers/:id, /interpreters", "agency_admin, scheduler"],
        ["/billing-rates, /invoices, /audit-log, /import, etc.", "agency_admin only"],
        ["/request, /my-requests", "requester"],
        ["/my-schedule, /availability, /available-jobs, etc.", "interpreter"],
    ])
    add_para(doc, "Platform routes use separate PlatformGuard checking platform_roles.")

    # Part D
    add_heading(doc, "Part D — Appendix: Verification Checklist (Operator)", 1)
    checklist = [
        "npm install && npm run build — no errors",
        "Supabase: 100 migrations applied — no drift",
        "Edge Functions deployed — 13 functions",
        "Secrets: RESEND, TWILIO, GOOGLE, QBO, APP_BASE_URL",
        "VITE_* set in frontend build environment",
        "Login as each role — smoke test",
        "Requester cannot read other customers' appointments (API test)",
        "process-reminders rejects unauthenticated POST",
        "/demo returns 404 or redirect in production build",
        "QBO webhook signature test in sandbox",
        "Google Calendar round-trip on test appointment",
        "npm audit — review critical CVEs",
        "Legal: BAAs executed for Supabase, Resend, Twilio, Google, Intuit",
    ]
    for c in checklist:
        p = doc.add_paragraph()
        run = p.add_run(f"☐ {c}")
        run.font.size = Pt(11)

    doc.add_page_break()

    add_heading(doc, "5. Recommended Next Steps (Numbered)", 1)
    steps = [
        "Obtain Supabase project URL/keys (or create staging) and confirm all 100 migrations are applied without drift.",
        "Run locally: npm install, configure .env from .env.example, npm run dev, npm run build, npm run lint.",
        "Deploy all 13 Edge Functions and configure secrets: RESEND_API_KEY, TWILIO_*, GOOGLE_CALENDAR_*, QBO_*, APP_BASE_URL, QBO_WEBHOOK_VERIFIER_TOKEN.",
        "Execute role-based security test plan on staging (admin, scheduler, requester, interpreter, cross-tenant negative tests).",
        "Decide with client: QBO-only agency invoicing vs Stripe card payments (new development if cards required).",
        "Decide HIPAA scope; if yes, engage compliance advisor and execute vendor BAAs (Supabase, Resend, Twilio, Google, Intuit, Lovable/hosting).",
        "Gate demo mode off for production build; remove or hide /demo route.",
        "Fix process-reminders authentication before exposing production Supabase function URL.",
        "Run npm audit and address critical dependency CVEs; add Vitest and lint to CI pipeline.",
        "Complete UAT using Playwright E2E with .env.test credentials against staging.",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"Step {i}: {s}")

    add_heading(doc, "6. Client Questions Before Final Scope (Summary)", 1)
    add_para(doc, "See Section 19 for full list. Minimum questions before quoting: regulatory scope (HIPAA?), payment type (QBO vs cards), hosting commitment (Supabase/Lovable vs AWS), calendar providers (Google/Outlook), demo mode policy, MFA/SSO requirements, existing live Supabase project with data, and data residency.")

    doc.add_page_break()

    # ========== APPENDIX E: Backend modules checklist ==========
    add_heading(doc, "Appendix E — Required Backend Modules Checklist", 1)
    add_para(doc, "Assessment of modules from the original audit scope. Status reflects codebase evidence only.")
    add_table(doc, ["Module", "Status", "Evidence"], [
        ["Authentication/session handling", "Partial", "Supabase Auth + AuthContext.tsx; no MFA"],
        ["Role-based access control", "Partial", "user_roles + RLS + route-roles.ts (UI); pen-test required"],
        ["User profile management", "Partial", "profiles table + Settings; agency_id immutability fixed in migration"],
        ["Interpreter profile management", "Partial", "Interpreters.tsx, interpreter_languages, invite-user"],
        ["Client/customer management", "Partial", "customers, locations, customer_requestors"],
        ["Booking/scheduling engine", "Partial", "appointments + Schedule Wizard + search_appointments RPC"],
        ["Availability management", "Partial", "interpreter_availability with recurring + specific_date"],
        ["Calendar integration", "Partial", "google-calendar-sync; Outlook not found"],
        ["Payment processing (card)", "Missing", "Stripe/PayPal not found"],
        ["Payment webhook handling (QBO)", "Partial", "qbo-webhook with HMAC verification"],
        ["Notification system", "Partial", "send-notification, process-reminders, Resend, Twilio"],
        ["Audit logging", "Partial", "appointment_history, platform_audit_log; not full PHI access log"],
        ["Admin management", "Partial", "agency_admin routes + platform-admin Edge Function"],
        ["File/document upload", "Missing", "No supabase.storage.from patterns found in src"],
        ["HIPAA-conscious data handling", "Low", "patient_client_name field; no technical control package"],
    ])

    add_heading(doc, "Appendix F — Recommended Safe Payment Architecture", 1)
    add_para(doc, "If the client requires online card payments at launch (not currently in codebase):")
    add_bullet(doc, "Use Stripe Checkout or Payment Element — card data never touches this React app or Supabase tables directly.")
    add_bullet(doc, "Create payment intent/session via new Edge Function using STRIPE_SECRET_KEY (server-only).")
    add_bullet(doc, "Handle stripe webhooks in dedicated Edge Function with signature verification (similar pattern to qbo-webhook).")
    add_bullet(doc, "Store only Stripe customer ID, payment intent ID, and payment status in invoices or new payments table.")
    add_bullet(doc, "Never place sk_live or secret keys in Vite frontend (VITE_* variables are public in bundle).")
    add_para(doc, "If QBO-only (current design): keep card data out of app; reconcile invoice status via QBO webhooks and manual/admin workflows.")

    add_heading(doc, "Appendix G — Scheduling Edge Cases", 1)
    add_table(doc, ["Edge Case", "Status in Codebase"], [
        ["Double booking same interpreter", "Client-side warning + admin override only; no DB exclusion constraint"],
        ["Daylight saving time", "agency-timezone.ts + Schedule Wizard comments address DST in availability walk"],
        ["Recurring appointments", "parent_recurring_id, bulkCreate/bulkUpdate in useAgencyData.ts"],
        ["Cancellation window / fees", "billing-engine cancellation_window_hours, cancellation_fee_percent"],
        ["No-show handling", "Statuses: no_show_interpreter, late_cancel_no_show_client in import validator"],
        ["Time zone per agency", "agencies.timezone; resolveTimezone in google-calendar-sync"],
        ["Conflict override audit", "override_conflict in appointment_history + custom_fields.override_log"],
        ["Self-claimable jobs", "is_self_claimable on appointments; AvailableJobs page"],
        ["Last-minute requests", "requested_last_minute, interpreter_assigned_last_minute statuses"],
        ["Import-staged appointments", "is_import_staged filter on queries"],
    ])

    add_heading(doc, "Appendix H — Output Format Index (Per Original Audit Request)", 1)
    add_para(doc, "This document is organized to satisfy the requested deliverables:")
    add_table(doc, ["#", "Deliverable", "Location in Document"], [
        ["1", "Internal Developer Notes", "Part A (sections A.1–A.4)"],
        ["2", "Client-Ready Audit Report", "Part B (sections 1–20, Executive Summary, G, C, D, E, F)"],
        ["3", "Feature Completion Matrix", "Section 15"],
        ["4", "Security Risk Matrix", "Section 16"],
        ["5", "Recommended Next Steps", "Section 17, A.4, Appendix verification checklist"],
        ["6", "Client Questions Before Final Scope", "Section 19"],
    ])

    add_heading(doc, "Appendix I — supabase/config.toml Edge Function JWT Settings", 1)
    add_para(doc, "All functions listed below have verify_jwt = false. Custom authentication must be verified per function.")
    for fn in [
        "auth-email-hook", "google-calendar-sync", "send-notification", "process-reminders",
        "process-import", "qbo-auth", "qbo-webhook", "platform-admin", "platform-qbo",
        "invite-user", "manage-join-request",
    ]:
        add_bullet(doc, fn)

    add_heading(doc, "Appendix J — Playwright E2E Test Files", 1)
    for spec in [
        "e2e/auth.spec.ts — login, signup validation, unauthenticated redirect",
        "e2e/appointments.spec.ts — appointment CRUD UI (requires admin credentials)",
        "e2e/billing.spec.ts — billing flows",
        "e2e/customers.spec.ts — customer management",
        "e2e/demo.spec.ts — demo mode entry/exit (no Supabase required)",
        "e2e/interpreter.spec.ts — interpreter role flows",
        "e2e/messaging.spec.ts — messaging UI",
        "e2e/requester.spec.ts — requester role flows",
        "e2e/responsive.spec.ts — responsive layout",
        "e2e/settings-notifications.spec.ts — settings and notifications",
        "e2e/admin-navigation.spec.ts — admin navigation",
    ]:
        add_bullet(doc, spec)

    add_heading(doc, "Appendix K — Data-Adapter Pattern (Technical Reference)", 1)
    add_para(doc, "File: src/lib/data-adapter.ts")
    add_para(doc, "useAdaptedQuery: When isDemoMode is true, returns demoFn() result synchronously with isLoading=false without calling Supabase. When false, standard useQuery runs queryFn against Supabase.")
    add_para(doc, "useAdaptedMutation: mutationFn branches on isDemoMode — demoFn updates DemoDataContext in memory; production mutationFn hits Supabase and invalidates query keys on success.")
    add_para(doc, "Convention documented in file header: demoFn/demoMutationFn operate on DemoDataContext; queryFn/mutationFn hit Supabase.")

    add_heading(doc, "Appendix L — process-reminders Authentication Logic (Verbatim Summary)", 1)
    add_para(doc, "File: supabase/functions/process-reminders/index.ts")
    add_bullet(doc, "If Authorization header present and token equals SUPABASE_SERVICE_ROLE_KEY: allow.")
    add_bullet(doc, "Else if Authorization header present: verify JWT via getClaims; require agency_admin role in user_roles.")
    add_bullet(doc, "Else (no Authorization header): allow — documented as backward-compatible cron invocations.")
    add_para(doc, "Production risk: public URL + verify_jwt=false + no header = unauthenticated execution with service-role DB access inside function body.", bold=True)

    # Footer note
    doc.add_paragraph()
    add_para(doc, "End of expanded audit. All file paths are relative to repository root bluethreadsolutions-main. Findings are based on static analysis; live system behavior may differ if the deployed Supabase project has diverged from migration history.", italic=True)
    add_para(doc, "Document generated by generate_audit_docx.py from static codebase audit performed May 23, 2026.", italic=True)

    return doc


def main():
    import os
    script_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(script_dir, OUTPUT)
    doc = build_document()
    doc.save(out_path)
    print(f"Written: {out_path}")
    print(f"Size: {os.path.getsize(out_path)} bytes")


if __name__ == "__main__":
    main()
